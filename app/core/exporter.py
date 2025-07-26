from docx import Document
from docx.shared import Inches
import PyPDF2
from tabulate import tabulate
import json
import csv
from datetime import datetime

class ResearchExporter:
    def __init__(self):
        self.document = None

    def create_document(self):
        self.document = Document()
        self.document.add_heading('Research Report', 0)
        self.document.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    def add_section(self, title, content):
        self.document.add_heading(title, level=1)
        self.document.add_paragraph(content)

    def add_table(self, data, headers):
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Add data
        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)

    def export_docx(self, filename):
        self.document.save(filename)

    def export_pdf(self, input_docx, output_pdf):
        # Note: This is a placeholder. In a real implementation,
        # you would need to use a PDF conversion library like reportlab
        # or a system command to convert DOCX to PDF
        pass

    def export_csv(self, data, filename, headers=None):
        with open(filename, 'w', newline='') as f:
            writer = csv.writer(f)
            if headers:
                writer.writerow(headers)
            writer.writerows(data)

    def export_json(self, data, filename):
        with open(filename, 'w') as f:
            json.dump(data, f, indent=2) 